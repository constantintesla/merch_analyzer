"""Заполнение Приложения №1 (трудозатраты) в акте Маскон — апрель 2026."""
from __future__ import annotations

from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

DOC_PATH = "ИП Тюменцев_Маскон_Акт приемки услуг_Апрель 2026.docx"

# Ставка 1250 ₽/ч. Базовая оценка 110 ч → 137 500 ₽; +10 % на налоги/накладные → 151 250 ₽ (121 ч).
# Часы по строкам: базовые 14/36/32/18/10 равномерно увеличены на 10 %.
HOURLY_RATE = 1250
TOTAL_RUB = 151_250
TOTAL_HOURS = TOTAL_RUB // HOURLY_RATE  # 121

ROWS: list[tuple[str, str, str, str]] = [
    (
        "1",
        "Проанализировали постановку задачи заказчика, сформулировали и согласовали пайплайн "
        "обработки фото торговой полки; настроили среду разработки (Python, зависимости), "
        "развернули репозиторий и файл весов SKU110K",
        "15,4",
        "19250.00",
    ),
    (
        "2",
        "Разработали серверное приложение на FastAPI: реализовали приём эталонного снимка "
        "и детекцию позиций, настроили сохранение прогонов, спроектировали эндпоинты "
        "распознавания и выдачу артефактов (JSON, изображения, кропы)",
        "39,6",
        "49500.00",
    ),
    (
        "3",
        "Интегрировали детектор SKU110K/RetinaNet: написали адаптер вызова (Docker/WSL), "
        "реализовали нарезку и сохранение кропов, настроили оценку рядов полок и экспорт разметки",
        "35,2",
        "44000.00",
    ),
    (
        "4",
        "Разработали HTTP-клиент к LM Studio (vision), настроили пакетную отправку кропов "
        "и снизили нагрузку на модель за счёт группировки похожих изображений",
        "19,8",
        "24750.00",
    ),
    (
        "5",
        "Доработали веб-шаблон и визуализацию боксов и полок; провели отладку на данных "
        "заказчика, прогнали сценарии тестирования, актуализировали README по развёртыванию",
        "11",
        "13750.00",
    ),
]


def _insert_data_rows_before_total(table, n_extra: int) -> None:
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    if len(trs) < 2:
        raise RuntimeError("Таблица слишком мала")
    template = trs[1]
    total_tr = trs[-1]
    for _ in range(n_extra):
        total_tr.addprevious(deepcopy(template))


def _replace_in_paragraph(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    text = paragraph.text.replace(old, new)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _sync_act_amounts_in_body(doc: Document) -> None:
    """Сумма в тексте акта: синхронизация с таблицей (в т.ч. после смены базы или +10 %)."""
    amount_new = "151 250,00"
    spelling_new = "Сто пятьдесят одна тысяча двести пятьдесят рублей"
    amount_spell_pairs = [
        ("164 000,00", "Сто шестьдесят четыре тысячи рублей"),
        ("137 500,00", "Сто тридцать семь тысяч пятьсот рублей"),
    ]

    for p in doc.paragraphs:
        for amount_old, spelling_old in amount_spell_pairs:
            if amount_old in p.text:
                _replace_in_paragraph(p, amount_old, amount_new)
            if spelling_old in p.text:
                _replace_in_paragraph(p, spelling_old, spelling_new)


def main() -> None:
    doc = Document(DOC_PATH)
    labor = doc.tables[1]
    trs = labor._tbl.findall(qn("w:tr"))
    if len(trs) == 4:
        _insert_data_rows_before_total(labor, 3)
        trs = labor._tbl.findall(qn("w:tr"))
    if len(trs) != 7:
        raise RuntimeError(f"Ожидалось 7 строк таблицы трудозатрат, получилось {len(trs)}")

    for i, (num, title, hours, cost) in enumerate(ROWS, start=1):
        row = labor.rows[i]
        row.cells[0].text = num
        row.cells[1].text = title
        row.cells[2].text = hours
        row.cells[3].text = cost

    total_row = labor.rows[6]
    total_row.cells[0].text = "Итого"
    total_row.cells[2].text = str(TOTAL_HOURS)
    total_row.cells[3].text = f"{TOTAL_RUB}.00"

    bad = "ИП Барков Юрий Александрович"
    good = "ИП Тюменцев Павел Андреевич"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if bad in cell.text:
                    cell.text = cell.text.replace(bad, good)

    _sync_act_amounts_in_body(doc)

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
